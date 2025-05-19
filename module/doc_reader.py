import docx
from striprtf.striprtf import rtf_to_text
import os
import re # Для некоторых проверок
from reportlab.lib.pagesizes import A4

from module.config import id2label

def classify_docx_paragraph(paragraph):
    """
    Пытается классифицировать параграф DOCX.
    Возвращает (type_id, type_label).
    """
    style_name = paragraph.style.name.lower()
    text = paragraph.text.strip()

    if not text: # Пустые параграфы пропускаем или считаем текстом? Пока пропустим.
        return None, None

    # Заголовки
    if style_name.startswith('heading 1'):
        return "0", id2label["0"]
    if style_name.startswith('heading'): # Heading 2, Heading 3 и т.д.
        return "7", id2label["7"]

    
    if paragraph.style.name == 'List Paragraph':
        return "3", id2label["3"]
    # Более сложная проверка на маркеры или нумерацию (требует анализа run.text)
    # Например, если первый run содержит '•', '*', или цифру с точкой.
    # Этот подход очень упрощен и может давать ложные срабатывания.
    first_run_text = paragraph.runs[0].text.strip() if paragraph.runs else ""
    if re.match(r"^(\d+\.|[a-zA-Z]\.|[•*-])", first_run_text):
         if len(text) > len(first_run_text) + 2: # Чтобы не принять просто "1." за список
            return "3", id2label["3"]


    # По умолчанию - обычный текст
    return "9", id2label["9"]

def extract_from_docx(filepath):
    """
    Извлекает текст, "позицию" (описательную) и тип из DOCX.
    """
    doc = docx.Document(filepath)
    extracted_data = []
    element_index = 0
    len_element = len(doc.element.body)
    # 1. Основной текст документа (параграфы и таблицы)
    for element in doc.element.body:
        element_index += 1
        if isinstance(element, docx.oxml.text.paragraph.CT_P): # Параграф
            para = docx.text.paragraph.Paragraph(element, doc)
            text = para.text.strip()
            if not text:
                continue

            type_id, type_label = classify_docx_paragraph(para)
            if type_id is None: # Пропускаем если классификатор вернул None
                continue

            bbox = (4,element_index*(A4[1]/len_element), A4[0]-4, (element_index+1)*(A4[1]/len_element))
            extracted_data.append({
                "text": text,
                "bbox": bbox,
                "type_id": type_id,
                "type_label": type_label
            })

        elif isinstance(element, docx.oxml.table.CT_Tbl): # Таблица
            table = docx.table.Table(element, doc)
            table_text_parts = []
            for i, row in enumerate(table.rows):
                row_texts = []
                for j, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    row_texts.append(cell_text)
                   
                table_text_parts.append(" | ".join(row_texts)) # Соединяем ячейки строки
            full_table_text = "\n".join(table_text_parts) # Соединяем строки
            bbox = (4,element_index*(A4[1]/len_element), A4[0]-4, (element_index+1)*(A4[1]/len_element))

            if full_table_text:
                 extracted_data.append({
                    "text": full_table_text,
                    "bbox": bbox,
                    "type_id": "8",
                    "type_label": id2label["8"]
                })

    # 2. Верхние и нижние колонтитулы
    for section in doc.sections:
        # Верхний колонтитул
        for para in section.header.paragraphs:
            text = para.text.strip()    
            if text:
                element_index += 1
                bbox = (4,element_index*(A4[1]/len_element), A4[0]-4, (element_index+1)*(A4[1]/len_element))
                extracted_data.append({
                    "text": text,
                    "bbox": bbox,
                    "type_id": "5",
                    "type_label": id2label["5"]
                })
        # Нижний колонтитул
        for para in section.footer.paragraphs:
            text = para.text.strip()
            if text:
                element_index += 1
                bbox = (4,element_index*(A4[1]/len_element), A4[0]-4, (element_index+1)*(A4[1]/len_element))
                extracted_data.append({
                    "text": text,
                    "bbox": bbox,
                    "type_id": "4",
                    "type_label": id2label["4"]
                })

    data_from_docx = {
        "texts": [i["text"] for i in extracted_data],
        "bboxes": [i["bbox"] for i in extracted_data],
        "labels": [i["type_label"] for i in extracted_data]
    }

    return data_from_docx

# --- RTF Parsing ---

def extract_from_rtf(filepath):
    """
    Извлекает текст из RTF. Определение типа и позиции очень ограничено.
    striprtf просто извлекает весь текст как единый блок.
    """
    extracted_data = []
    try:
        with open(filepath, 'r', encoding='utf-8') as f: # Пытаемся UTF-8
            rtf_content = f.read()
    except UnicodeDecodeError:
        try:
            with open(filepath, 'r', encoding='latin-1') as f: # Частая кодировка для RTF
                rtf_content = f.read()
        except Exception as e:
            print(f"Не удалось прочитать RTF файл {filepath}: {e}")
            return []

    text = rtf_to_text(rtf_content).strip()

    if text:
        len_spilt = len(text.split("\n"))
        for e,i in enumerate(text.split("\n")):
            bbox = (4,e*(A4[1]/len_spilt), A4[0], (e+1)*(A4[1]/len_spilt))
            extracted_data.append({
                "text": i,
                "bbox": bbox,
                "type_id": "9",
                "type_label": id2label["9"]
            })
    data_from_docx = {
        "texts": [i["text"] for i in extracted_data],
        "bboxes": [i["bbox"] for i in extracted_data],
        "labels": [i["type_label"] for i in extracted_data]
    }
    return data_from_docx

def put_to_docx(filepath, extracted_data):
    """
    Добавляет текст в DOCX.
    """
    doc = docx.Document(filepath)
    texts = extracted_data.get('rephrased_texts', extracted_data.get('texts', []))
    
    # Make sure we have texts to add
    if not texts:
        print(f"Warning: No texts found in extracted_data for file {filepath}")
        return doc
    
    element_index = 0
    text_index = 0
    max_text_index = len(texts) - 1
    
    # 1. Основной текст документа (параграфы и таблицы)
    for element in doc.element.body:
        if text_index > max_text_index:
            # No more texts to add
            break
            
        if isinstance(element, docx.oxml.text.paragraph.CT_P): # Параграф
            para = docx.text.paragraph.Paragraph(element, doc)
            if para.text.strip():  # Only replace non-empty paragraphs
                para.text = texts[text_index]
                text_index += 1
            
        elif isinstance(element, docx.oxml.table.CT_Tbl): # Таблица
            print("Table found, skipping")

    # 2. Верхние и нижние колонтитулы
    for section in doc.sections:
        # Верхний колонтитул
        for para in section.header.paragraphs:
            if text_index <= max_text_index and para.text.strip():
                para.text = texts[text_index]
                text_index += 1
                
        # Нижний колонтитул
        for para in section.footer.paragraphs:
            if text_index <= max_text_index and para.text.strip():
                para.text = texts[text_index]
                text_index += 1

    return doc

# --- Main Extractor ---

def extract_document_data(filepath):
    """
    Главная функция для извлечения данных из файла.
    """
    filename, file_extension = os.path.splitext(filepath)
    file_extension = file_extension.lower()

    if not os.path.exists(filepath):
        print(f"Файл не найден: {filepath}")
        return []

    print(f"Обработка файла: {filepath}")

    if file_extension == '.docx':
        return extract_from_docx(filepath)
    elif file_extension == '.rtf':
        return extract_from_rtf(filepath)
    else:
        print(f"Неподдерживаемый формат файла: {file_extension}")
        return []
